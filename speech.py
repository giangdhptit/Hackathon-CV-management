import speech_recognition as sr
from pdf_generation import *

# def recognize_speech_from_mic():
#     recognizer = sr.Recognizer()
#     with sr.Microphone() as source:
#         print("Say something...")
#         recognizer.adjust_for_ambient_noise(source)  # Reduce background noise
#         audio = recognizer.listen(source)

#     try:
#         text = recognizer.recognize_google(audio)  # Use Google Speech API
#         print("You said:", text)
#         return text
#     except sr.UnknownValueError:
#         print("Sorry, could not understand the audio.")
#     except sr.RequestError:
#         print("Could not request results. Check your internet connection.")

# text = recognize_speech_from_mic()
# output_pdf = "/Users/giangdinh/Documents/projects/hackathon/output.pdf"
# create_pdf_from_text(text, output_pdf)


import speech_recognition as sr
import pyttsx3
from docx import Document

# Function to ask questions using text-to-speech (pyttsx3)
def ask_question(question):
    engine = pyttsx3.init()
    rate = engine.getProperty('rate')
    
    # Set the speech rate to a slower value (adjust as needed)
    engine.setProperty('rate', rate - 50)
    engine.say(question)
    engine.runAndWait()

# Function to capture audio input from the user (using speech_recognition)
# def get_user_input(prompt):
#     engine = pyttsx3.init()
#     rate = engine.getProperty('rate')
    
#     # Set the speech rate to a slower value (adjust as needed)
#     engine.setProperty('rate', rate - 50)
#     recognizer = sr.Recognizer()
#     while 1:
#         with sr.Microphone() as source:
#             print(prompt)
#             recognizer.adjust_for_ambient_noise(source)  # Adjust for ambient noise
#             audio = recognizer.listen(source)  # Listen to the user's answer
#         try:
#             user_confirmation = ""
#             user_input = ""
#             while (user_confirmation!="Yes"):
#                 user_input = recognizer.recognize_google(audio)  # Recognize the speech using Google API
#                 print(f"You said: {user_input}, is that correct?")
#                 engine.say("You said: {user_input}, is that correct?")
#                 audio = recognizer.listen(source)  # Listen to the user's answer
#                 user_confirmation = recognizer.recognize_google(audio)
#             return user_input
#         except sr.UnknownValueError:
#             print("Sorry, I could not understand that. Can you say it again?")
#             engine.say("Sorry, I could not understand that. Can you say it again?")
#             return ""
#         except sr.RequestError:
#             print("Sorry, I'm having trouble with the speech recognition service.")
#             return ""


recognizer = sr.Recognizer()
engine = pyttsx3.init()

import speech_recognition as sr
import pyttsx3

# Initialize recognizer and pyttsx3 engine
recognizer = sr.Recognizer()
engine = pyttsx3.init()

# Function to ask questions using text-to-speech (pyttsx3)
def ask_question(question):
    engine.say(question)
    engine.runAndWait()

import speech_recognition as sr
import pyttsx3

# Initialize recognizer and pyttsx3 engine
recognizer = sr.Recognizer()
engine = pyttsx3.init()

# Function to ask questions using text-to-speech (pyttsx3)
def ask_question(question):
    engine.say(question)
    engine.runAndWait()

# Function to capture audio input from the user (using speech_recognition)
def get_user_input(prompt):
    # Ensure the microphone source is inside the 'with' statement
    while True:
        with sr.Microphone() as source:
            print(prompt)
            recognizer.adjust_for_ambient_noise(source)  # Adjust for ambient noise
            audio = recognizer.listen(source)  # Listen to the user's answer
        
        user_input = ""
        user_confirmation = ""
        try:
            # Recognize speech
            user_input = recognizer.recognize_google(audio)
            print(f"You said: {user_input}, is that correct?")
            
            # Confirm the input with the user using TTS
            engine.say(f"You said: {user_input}, is that correct?")
            engine.runAndWait()
            
            # Listen for user confirmation (inside the same `with` statement)
            with sr.Microphone() as source:
                audio = recognizer.listen(source)
                user_confirmation = recognizer.recognize_google(audio)
            
            # If user confirms, break the loop and return input
            if user_confirmation.lower() in ["yes", "y"]:
                return user_input
            else:
                # If the user says something other than "Yes", prompt again
                print("Please repeat your response.")
                engine.say("Please repeat your response.")
                engine.runAndWait()
                continue  # Retry the input

        except sr.UnknownValueError:
            print("Sorry, I could not understand that. Can you say it again?")
            engine.say("Sorry, I could not understand that. Can you say it again?")
            engine.runAndWait()
            continue  # Retry on recognition failure
        except sr.RequestError:
            print("Sorry, I'm having trouble with the speech recognition service.")
            engine.say("Sorry, I'm having trouble with the speech recognition service.")
            engine.runAndWait()
            return ""  # Retry on request failure

# Function to generate a docx file with sections and user responses
def generate_docx(data, output_filename):
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)

    for section, content in data.items():
        doc.add_heading(section, level=1)
        for item in content:
            doc.add_paragraph(f"- {item}")

    doc.save(output_filename)
    print(f"Document saved as {output_filename}")

# Main process to interact with the user and collect the information
def create_cv():
    sections = {
        "Name": "What is your full name?",
        "Email": "What is your email address?",
        "Skills": "What are your key skills? Please list them.",
        "Experience": "Please describe your work experience.",
        "Education": "Where did you study and what degree did you earn?",
        "Certification": "Do you have any certifications or awards?",
        "References": "Please provide any references."
    }

    # Collect data
    cv_data = {}

    for section, question in sections.items():
        ask_question(question)  # Ask the user the question
        response = get_user_input(f"Please answer the question: {question}")  # Get the response from the user
        cv_data[section] = response.split(",")  # Split by commas if multiple items (for example, skills)

    # Generate DOCX
    output_filename = "/Users/giangdinh/Documents/projects/hackathon/output.docx"
    generate_docx(cv_data, output_filename)

# Run the create_cv function to start the process
create_cv()
